"""
utils/doc_generator.py
======================
Yo'qolgan tovarlar uchun hujjat generatsiya:
  1. generate_xlsx()  → openpyxl  — отчет-потерянных-товаров.xlsx
  2. generate_docx()  → python-docx — Доп-соглашение.docx
  3. create_zip()     → ikki faylni bitta .zip ga joylashtirish
"""
import io
import logging
import zipfile
from datetime import datetime

logger = logging.getLogger(__name__)


# ─── Excel ────────────────────────────────────────────────────────────────────

def generate_xlsx(form: dict, items: list) -> io.BytesIO:
    """
    отчет-потерянных-товаров.xlsx ni generatsiya qiladi.

    form keys:
        report_date, director_name, org_name, reg_number, inn,
        address, contract_number, contract_date,
        bank_name, account_number, bank_mfo
    items: list[MissingItem]
    """
    try:
        import openpyxl
        from openpyxl.styles import (
            Font, PatternFill, Alignment, Border, Side
        )
        from openpyxl.utils import get_column_letter
    except ImportError as e:
        logger.error(f"openpyxl yo'q: {e}")
        return io.BytesIO(b"")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Yo'qolgan tovarlar"

    # ── rang palitasi ──────────────────────────────────────────────────────────
    BLUE  = "1F497D"
    LBLUE = "DCE6F1"
    GRAY  = "808080"
    RED   = "C00000"
    WHITE = "FFFFFF"
    LGRAY = "F2F2F2"

    def cell(r, c, val="", bold=False, size=10, color=None,
             bg=None, align="left", wrap=False, number_format=None):
        cc = ws.cell(row=r, column=c, value=val)
        cc.font = Font(bold=bold, size=size,
                       color=color or "000000",
                       name="Calibri")
        if bg:
            cc.fill = PatternFill("solid", fgColor=bg)
        cc.alignment = Alignment(
            horizontal=align, vertical="center",
            wrap_text=wrap
        )
        if number_format:
            cc.number_format = number_format
        return cc

    thin = Side(style="thin")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    def border_range(r1, c1, r2, c2):
        for row in ws.iter_rows(min_row=r1, max_row=r2,
                                min_col=c1, max_col=c2):
            for cc in row:
                cc.border = border

    # ── sarlavha ───────────────────────────────────────────────────────────────
    ws.merge_cells("A1:G1")
    cell(1, 1,
         "ОТЧЁТ О ПОТЕРЯННЫХ ТОВАРАХ НА СКЛАДЕ UZUM (FBO)",
         bold=True, size=13, color=WHITE, bg=BLUE,
         align="center")
    ws.row_dimensions[1].height = 28

    # ── kompaniya ma'lumotlari ─────────────────────────────────────────────────
    info_rows = [
        ("Дата отчёта",          form.get("report_date", "")),
        ("Руководитель (Ф.И.О)", form.get("director_name", "")),
        ("Наименование организации", form.get("org_name", "")),
        ("Рег. номер (ЯТТ/ОГРНИП)", form.get("reg_number", "")),
        ("ИНН / ЖШШИР",          form.get("inn", "")),
        ("Адрес",                form.get("address", "")),
        ("Номер договора Uzum",  form.get("contract_number", "")),
        ("Дата договора",        form.get("contract_date", "")),
    ]

    row = 3
    ws.merge_cells(f"A{row}:G{row}")
    cell(row, 1, "СВЕДЕНИЯ ОБ ОРГАНИЗАЦИИ", bold=True,
         size=11, color=WHITE, bg=BLUE, align="center")
    ws.row_dimensions[row].height = 20
    row += 1

    for label, value in info_rows:
        ws.merge_cells(f"A{row}:C{row}")
        cell(row, 1, label, bold=True, bg=LGRAY)
        ws.merge_cells(f"D{row}:G{row}")
        cell(row, 4, value)
        border_range(row, 1, row, 7)
        ws.row_dimensions[row].height = 18
        row += 1

    # ── bank rekvizitlari ──────────────────────────────────────────────────────
    row += 1
    ws.merge_cells(f"A{row}:G{row}")
    cell(row, 1, "БАНКОВСКИЕ РЕКВИЗИТЫ", bold=True,
         size=11, color=WHITE, bg=BLUE, align="center")
    ws.row_dimensions[row].height = 20
    row += 1

    bank_rows = [
        ("Банк",            form.get("bank_name", "")),
        ("Расчётный счёт",  form.get("account_number", "")),
        ("МФО",             form.get("bank_mfo", "")),
    ]
    for label, value in bank_rows:
        ws.merge_cells(f"A{row}:C{row}")
        cell(row, 1, label, bold=True, bg=LGRAY)
        ws.merge_cells(f"D{row}:G{row}")
        cell(row, 4, value)
        border_range(row, 1, row, 7)
        ws.row_dimensions[row].height = 18
        row += 1

    # ── yo'qolgan tovarlar jadvali ─────────────────────────────────────────────
    row += 1
    ws.merge_cells(f"A{row}:G{row}")
    cell(row, 1, "ПЕРЕЧЕНЬ ПОТЕРЯННЫХ ТОВАРОВ", bold=True,
         size=11, color=WHITE, bg=RED, align="center")
    ws.row_dimensions[row].height = 20
    row += 1

    # jadval boshi
    headers = ["№", "Наименование товара", "SKU / Код",
               "Штрих-код", "Кол-во (шт.)", "Цена (сум)", "Компенсация (сум)"]
    col_widths = [5, 40, 18, 18, 12, 15, 18]

    for c, (h, w) in enumerate(zip(headers, col_widths), 1):
        cell(row, c, h, bold=True, color=WHITE, bg=BLUE,
             align="center", wrap=True)
        ws.column_dimensions[get_column_letter(c)].width = w
    ws.row_dimensions[row].height = 22
    border_range(row, 1, row, 7)
    row += 1

    # jadval satrlari
    data_start = row
    total_qty = 0
    total_comp = 0.0

    for idx, item in enumerate(items, 1):
        bg = WHITE if idx % 2 else LGRAY
        cell(row, 1, idx, align="center", bg=bg)
        cell(row, 2, item.name, wrap=True, bg=bg)
        cell(row, 3, item.sku_code, align="center", bg=bg)
        cell(row, 4, item.barcode, align="center", bg=bg)
        cell(row, 5, item.missing_qty, align="center", bold=True, bg=bg)
        cell(row, 6, item.purchase_price, align="right", bg=bg,
             number_format='#,##0.00')
        cell(row, 7, item.compensation, align="right", bold=True, bg=bg,
             number_format='#,##0.00')
        border_range(row, 1, row, 7)
        ws.row_dimensions[row].height = 20
        total_qty  += item.missing_qty
        total_comp += item.compensation
        row += 1

    # jami satr
    ws.merge_cells(f"A{row}:D{row}")
    cell(row, 1, "ИТОГО:", bold=True, color=WHITE, bg=RED, align="right")
    cell(row, 5, total_qty,  bold=True, color=WHITE, bg=RED,
         align="center", number_format='#,##0')
    cell(row, 6, "",          bg=RED)
    cell(row, 7, total_comp, bold=True, color=WHITE, bg=RED,
         align="right", number_format='#,##0.00')
    border_range(row, 1, row, 7)
    ws.row_dimensions[row].height = 22
    row += 2

    # ── imzo qatori ────────────────────────────────────────────────────────────
    ws.merge_cells(f"A{row}:C{row}")
    cell(row, 1, "Руководитель:", bold=True)
    ws.merge_cells(f"D{row}:G{row}")
    cell(row, 4, f"_____________ / {form.get('director_name', '')} /")
    ws.row_dimensions[row].height = 20
    row += 1
    ws.merge_cells(f"A{row}:C{row}")
    cell(row, 1, "Дата составления:", bold=True)
    ws.merge_cells(f"D{row}:G{row}")
    cell(row, 4, form.get("report_date", datetime.now().strftime("%Y-%m-%d")))
    ws.row_dimensions[row].height = 18

    # print area
    ws.print_area = f"A1:G{row}"
    ws.page_setup.fitToPage = True
    ws.page_setup.fitToWidth = 1

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


# ─── Word ─────────────────────────────────────────────────────────────────────

def generate_docx(form: dict, items: list) -> io.BytesIO:
    """
    Дополнительное соглашение .docx ni generatsiya qiladi.
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError as e:
        logger.error(f"python-docx yo'q: {e}")
        return io.BytesIO(b"")

    doc = Document()

    # sahifa chegaralari
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(1.5)

    def add_heading(text, level=1, center=False):
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14 if level == 1 else 12)
        run.font.name = "Times New Roman"
        if center:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p

    def add_para(text, bold_label=None, indent=False):
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        if bold_label:
            r = p.add_run(bold_label + " ")
            r.bold = True
            r.font.size = Pt(11)
            r.font.name = "Times New Roman"
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
        r2.font.name = "Times New Roman"
        if indent:
            p.paragraph_format.first_line_indent = Cm(1.25)
        return p

    def set_col_width(table, col_idx, width_cm):
        for row in table.rows:
            row.cells[col_idx].width = Cm(width_cm)

    # ── sarlavha ───────────────────────────────────────────────────────────────
    add_heading("ДОПОЛНИТЕЛЬНОЕ СОГЛАШЕНИЕ", level=1, center=True)
    add_heading(
        f"о компенсации потерянных товаров на складе Uzum Market (FBO)",
        level=2, center=True
    )
    doc.add_paragraph()

    # ── tomonlar ───────────────────────────────────────────────────────────────
    add_para(
        f"{form.get('report_date', datetime.now().strftime('%Y-%m-%d'))} г.",
        bold_label="Дата составления:"
    )
    add_para(
        f"Договор №{form.get('contract_number', '—')} "
        f"от {form.get('contract_date', '—')}",
        bold_label="Основание:"
    )
    doc.add_paragraph()

    add_para("Продавец (Сторона 1):", bold_label="")
    tbl_info = doc.add_table(rows=6, cols=2)
    tbl_info.style = "Table Grid"
    tbl_info.alignment = WD_TABLE_ALIGNMENT.LEFT

    info_data = [
        ("Наименование организации", form.get("org_name", "")),
        ("Ф.И.О. руководителя",      form.get("director_name", "")),
        ("Рег. номер (ЯТТ/ОГРНИП)",  form.get("reg_number", "")),
        ("ИНН / ЖШШИР",              form.get("inn", "")),
        ("Юридический адрес",        form.get("address", "")),
        ("Банковские реквизиты",
         f"Банк: {form.get('bank_name', '')} | "
         f"Счёт: {form.get('account_number', '')} | "
         f"МФО: {form.get('bank_mfo', '')}"),
    ]
    col_w = [6, 12]
    for i, (label, val) in enumerate(info_data):
        row_cells = tbl_info.rows[i].cells
        row_cells[0].text = label
        row_cells[1].text = val
        for ci, cw in enumerate(col_w):
            row_cells[ci].width = Cm(cw)
        # bold label
        for run in row_cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"
        for run in row_cells[1].paragraphs[0].runs:
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"

    doc.add_paragraph()

    # ── asosiy matn ────────────────────────────────────────────────────────────
    add_para(
        "В соответствии с условиями договора оферты, заключённого между Продавцом "
        "и Uzum Market, Продавец обнаружил расхождение между отправленными на склад "
        "товарами и фактическим наличием. Ниже приведён перечень потерянных товаров, "
        "по которым Продавец требует компенсацию.",
        indent=True
    )
    doc.add_paragraph()

    # ── jadval ─────────────────────────────────────────────────────────────────
    add_heading("Перечень потерянных товаров:", level=2)

    col_headers = ["№", "Наименование товара", "SKU / Код",
                   "Штрих-код", "Кол-во", "Цена (сум)", "Компенсация (сум)"]
    col_widths_cm = [1, 6, 3, 3, 1.5, 2.5, 3]

    tbl = doc.add_table(rows=1 + len(items) + 1, cols=7)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # header
    hdr = tbl.rows[0].cells
    for c, (h, w) in enumerate(zip(col_headers, col_widths_cm)):
        hdr[c].text = h
        hdr[c].width = Cm(w)
        p = hdr[c].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"
        # blue bg
        tc = hdr[c]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "1F497D")
        shd.set(qn("w:color"), "FFFFFF")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)

    total_qty  = 0
    total_comp = 0.0

    for r_idx, item in enumerate(items, 1):
        row_cells = tbl.rows[r_idx].cells
        values = [
            str(r_idx),
            item.name[:60],
            item.sku_code,
            item.barcode,
            str(item.missing_qty),
            f"{item.purchase_price:,.0f}",
            f"{item.compensation:,.0f}",
        ]
        for c, (val, w) in enumerate(zip(values, col_widths_cm)):
            row_cells[c].text = val
            row_cells[c].width = Cm(w)
            p = row_cells[c].paragraphs[0]
            if c in (4, 5, 6):
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = "Times New Roman"
        total_qty  += item.missing_qty
        total_comp += item.compensation

    # jami
    total_row = tbl.rows[-1].cells
    # birlashtiramiz (0-3)
    total_row[0].merge(total_row[3])
    total_row[0].text = "ИТОГО:"
    for run in total_row[0].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = "Times New Roman"
    total_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    total_row[4].text = str(total_qty)
    total_row[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in total_row[4].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)

    total_row[6].text = f"{total_comp:,.0f}"
    total_row[6].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in total_row[6].paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)

    doc.add_paragraph()

    # ── talab ──────────────────────────────────────────────────────────────────
    add_para(
        f"На основании вышеизложенного, Продавец требует выплатить компенсацию "
        f"в размере {total_comp:,.0f} (сум) за {total_qty} единиц потерянного товара.",
        indent=True
    )
    doc.add_paragraph()

    # ── imzolar ────────────────────────────────────────────────────────────────
    sign_tbl = doc.add_table(rows=3, cols=2)
    sign_tbl.style = "Table Grid"
    sign_data = [
        ("Продавец:", "Представитель Uzum Market:"),
        (f"ФИО: {form.get('director_name', '')} ________", "ФИО: ____________________"),
        (f"Дата: {form.get('report_date', '')}", "Дата: ___________________"),
    ]
    for i, (l, r) in enumerate(sign_data):
        sign_tbl.rows[i].cells[0].text = l
        sign_tbl.rows[i].cells[1].text = r
        for c in [0, 1]:
            sign_tbl.rows[i].cells[c].width = Cm(9)
            for run in sign_tbl.rows[i].cells[c].paragraphs[0].runs:
                run.font.size = Pt(10)
                run.font.name = "Times New Roman"
                if i == 0:
                    run.bold = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ─── ZIP ──────────────────────────────────────────────────────────────────────

def create_zip(form: dict, items: list) -> io.BytesIO:
    """
    Excel + Word → bitta .zip fayl.
    Qaytadi: io.BytesIO (ZIP ma'lumotlari).
    """
    report_date = form.get("report_date", datetime.now().strftime("%Y%m%d")).replace("-", "")

    xlsx_buf = generate_xlsx(form, items)
    docx_buf = generate_docx(form, items)

    zip_buf = io.BytesIO()
    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:
        if xlsx_buf.getbuffer().nbytes > 0:
            zf.writestr(
                f"отчет-потерянных-товаров-{report_date}.xlsx",
                xlsx_buf.getvalue()
            )
        if docx_buf.getbuffer().nbytes > 0:
            zf.writestr(
                f"Доп-соглашение-{report_date}.docx",
                docx_buf.getvalue()
            )

    zip_buf.seek(0)
    logger.info(f"ZIP yaratildi: {zip_buf.getbuffer().nbytes} bytes")
    return zip_buf
